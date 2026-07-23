"""从冻结的演示资料清单生成可由知识库解析器读取的多格式文件。"""

from __future__ import annotations

import argparse
import shutil
import sys
from pathlib import Path
from typing import Callable

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

# ``python scripts/generate_demo_knowledge_files.py`` places only the scripts
# directory on sys.path.  Add backend so the sibling ``scripts`` package can
# be imported both from the documented CLI and from pytest.
BACKEND_DIRECTORY = Path(__file__).resolve().parents[1]
if str(BACKEND_DIRECTORY) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIRECTORY))

from scripts.demo_knowledge_manifest import (
    KNOWLEDGE_DOCUMENTS,
    QUESTION_SETS,
    KnowledgeDocumentSpec,
    QuestionSpec,
)


KNOWLEDGE_DIRECTORY = "演示知识库资料"
QUESTIONS_DIRECTORY = "演示测试问题"
QUESTION_FILENAME = "测试问题与预期答案.md"
FONT_CANDIDATES = (
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\msyh.ttf"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
)


def _source_text(spec: KnowledgeDocumentSpec) -> str:
    return "\n\n".join((f"资料摘要：{spec.summary}", *spec.sections))


def _set_east_asia_font(run: object, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = font_name  # type: ignore[attr-defined]
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()  # type: ignore[attr-defined]
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:eastAsia"), font_name)


def _set_cell_shading(cell: object, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_margins(cell: object, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: object, widths: tuple[int, int] = (2700, 6660)) -> None:
    table.autofit = False  # type: ignore[attr-defined]
    tbl_pr = table._tbl.tblPr  # type: ignore[attr-defined]
    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    tbl_width.set(qn("w:w"), "9360")
    tbl_width.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    grid_columns = table._tbl.tblGrid.gridCol_lst  # type: ignore[attr-defined]
    for index, width in enumerate(widths):
        grid_columns[index].set(qn("w:w"), str(width))
    for row in table.rows:  # type: ignore[attr-defined]
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
            tc_width = tc_pr.first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(widths[index]))
            tc_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _add_page_field(paragraph: object) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # type: ignore[attr-defined]
    run = paragraph.add_run("第 ")  # type: ignore[attr-defined]
    _set_east_asia_font(run)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)  # type: ignore[attr-defined]
    run = paragraph.add_run(" 页")  # type: ignore[attr-defined]
    _set_east_asia_font(run)


def write_markdown(target: Path, spec: KnowledgeDocumentSpec) -> None:
    lines = [f"# {spec.title}", "", f"文档代号：{spec.document_code}", "", spec.summary]
    for index, section in enumerate(spec.sections, start=1):
        lines.extend(("", f"## 第 {index} 部分", "", section))
    target.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_text(target: Path, spec: KnowledgeDocumentSpec) -> None:
    lines = [spec.title, "=" * len(spec.title), f"文档代号：{spec.document_code}", "", spec.summary]
    for index, section in enumerate(spec.sections, start=1):
        lines.extend(("", f"第 {index} 部分", "-" * 12, section))
    target.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_docx(target: Path, spec: KnowledgeDocumentSpec) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.2)
    section.left_margin = section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = document.styles["Normal"]
    normal.font.name, normal.font.size = "Microsoft YaHei", Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = document.styles[name]
        style.font.name, style.font.size = "Microsoft YaHei", Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run(f"演示知识资料 | {spec.document_code}")
    _set_east_asia_font(header_run)
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(89, 89, 89)
    _add_page_field(section.footer.paragraphs[0])

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run(spec.title)
    _set_east_asia_font(title_run)
    title_run.font.size, title_run.bold = Pt(23), True
    title_run.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    code_run = subtitle.add_run(f"资料代号：{spec.document_code} | 资料类型：演示检索资料")
    _set_east_asia_font(code_run)
    code_run.font.size = Pt(10)
    code_run.font.color.rgb = RGBColor(89, 89, 89)

    document.add_heading("资料摘要", level=1)
    document.add_paragraph(spec.summary)
    document.add_heading("关键事实速览", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _set_table_geometry(table)
    headers = ("项目", "内容")
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        _set_cell_shading(cell, "E8EEF5")
        run = cell.paragraphs[0].add_run(value)
        _set_east_asia_font(run)
        run.bold = True
    for number, text in enumerate(spec.sections, start=1):
        row = table.add_row().cells
        left = row[0].paragraphs[0].add_run(f"第 {number} 部分")
        right = row[1].paragraphs[0].add_run(text)
        _set_east_asia_font(left)
        _set_east_asia_font(right)

    document.add_heading("详细说明", level=1)
    for number, text in enumerate(spec.sections, start=1):
        document.add_heading(f"第 {number} 部分", level=2)
        document.add_paragraph(text)
    document.save(target)


def _domain_rows(spec: KnowledgeDocumentSpec) -> list[tuple[str, str, str]]:
    rows: list[tuple[str, str, str]] = []
    for section_index, section in enumerate(spec.sections, start=1):
        for point_index in range(1, 5):
            excerpt_start = (point_index - 1) * max(1, len(section) // 4)
            excerpt = section[excerpt_start : excerpt_start + 130]
            rows.append((f"{spec.document_code}-{section_index}-{point_index}", f"第 {section_index} 部分", excerpt))
    return rows


def _style_sheet(sheet: object, widths: tuple[int, int, int] = (18, 22, 36)) -> None:
    blue = PatternFill("solid", fgColor="1F4D78")
    alternate = PatternFill("solid", fgColor="F2F4F7")
    thin = Side(style="thin", color="D9E2F3")
    for cell in sheet[1]:
        cell.fill = blue
        cell.font = Font(name="Microsoft YaHei", color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.font = Font(name="Microsoft YaHei", size=10)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(bottom=thin)
            if cell.row % 2 == 0:
                cell.fill = alternate
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    sheet.sheet_view.showGridLines = False
    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[get_column_letter(index)].width = width


def write_xlsx(target: Path, spec: KnowledgeDocumentSpec) -> None:
    workbook = Workbook()
    description = workbook.active
    description.title = "资料说明"
    description.append(("字段", "内容", "说明"))
    description.append(("文档标题", spec.title, "供知识库检索与引用验证使用"))
    description.append(("文档代号", spec.document_code, "每个资料唯一，回答引用时应保留"))
    description.append(("资料摘要", spec.summary, "全部为虚构教学数据，不含敏感信息"))
    description.append(("使用边界", "测试问题与预期答案不要上传知识库", "防止答案文本污染检索结果"))
    for index, section in enumerate(spec.sections, start=1):
        description.append((f"正文第 {index} 部分", section, "保留完整正文，供知识库按段检索"))
    _style_sheet(description, widths=(14, 36, 36))

    domain = workbook.create_sheet(spec.table_sheets[0] if spec.table_sheets else "资料明细")
    domain.append(("记录编号", "来源部分", "说明摘录"))
    for row in _domain_rows(spec):
        domain.append(row)
    _style_sheet(domain, widths=(26, 14, 36))
    workbook.save(target)


def _resolve_chinese_font() -> Path:
    for candidate in FONT_CANDIDATES:
        if candidate.exists():
            return candidate
    raise FileNotFoundError("没有找到可嵌入 PDF 的中文字体：msyh.ttc、msyh.ttf 或 simhei.ttf")


def _split_pdf_text(text: str) -> tuple[str, str]:
    """在中文句号处平衡切分两页，避免把整段文本固定堆到第二页。"""
    target = len(text) // 2
    split_at = text.find("。", target)
    if split_at == -1:
        split_at = text.rfind("。", 0, target)
    if split_at == -1:
        split_at = target
    else:
        split_at += 1
    return text[:split_at], text[split_at:]


def _pdf_pages(spec: KnowledgeDocumentSpec) -> tuple[tuple[str, str], tuple[str, str]]:
    content = "\n".join((spec.summary, *spec.sections))
    first_page, second_page = _split_pdf_text(content)
    return (("资料摘要与详细说明", first_page), ("详细说明（续）", second_page))


def write_pdf(target: Path, spec: KnowledgeDocumentSpec) -> None:
    font_path = _resolve_chinese_font()
    document = fitz.open()
    for page_number, (heading, text) in enumerate(_pdf_pages(spec), start=1):
        page = document.new_page(width=595.28, height=841.89)
        page.insert_font(fontname="CN", fontfile=str(font_path))
        page.insert_text((52, 52), f"演示知识资料 | {spec.document_code}", fontname="CN", fontsize=9, color=(0.35, 0.35, 0.35))
        y = 92
        if page_number == 1:
            page.insert_text((52, y), spec.title, fontname="CN", fontsize=22, color=(0.04, 0.15, 0.27))
            y += 34
            page.insert_text((52, y), f"资料代号：{spec.document_code}", fontname="CN", fontsize=10, color=(0.35, 0.35, 0.35))
            y += 34
        page.insert_text((52, y), heading, fontname="CN", fontsize=14, color=(0.18, 0.45, 0.71))
        y += 24
        remainder = page.insert_textbox(
            fitz.Rect(52, y, 543, 780), text, fontname="CN", fontsize=11.2,
            lineheight=1.5, color=(0.08, 0.08, 0.08),
        )
        if remainder < 0:
            raise ValueError(f"PDF 文本未能写入页面：{spec.filename}")
        page.insert_text((490, 806), f"第 {page_number} 页", fontname="CN", fontsize=9, color=(0.35, 0.35, 0.35))
    document.save(target, garbage=4, deflate=True)
    document.close()


def write_question_set(target: Path, folder: str, questions: tuple[QuestionSpec, ...]) -> None:
    lines = [
        f"# {folder}：测试问题与预期答案",
        "",
        "> **不要上传知识库。** 本文件含预期答案、评分标准和人工记录字段，仅用于演示验收。",
        "",
        "所有名称、规则、价格和日期均为虚构教学资料；回答必须以命中的知识文件与摘录为依据。",
    ]
    for number, question in enumerate(questions, start=1):
        sources = "、".join(question.source_files) if question.source_files else "无；应明确说明知识库缺少答案"
        points = "；".join(question.expected_points)
        lines.extend((
            "", f"### 问题 {number}", "", f"问题类型：{question.kind}", "", f"提问：{question.question}",
            "", f"预期要点：{points}", "", f"预期来源：{sources}", "", f"通过标准：{question.pass_criteria}",
            "", "实际回答：", "", "实际引用：", "", "验收结论：", "",
        ))
    target.write_text("\n".join(lines), encoding="utf-8")


WRITERS: dict[str, Callable[[Path, KnowledgeDocumentSpec], None]] = {
    ".md": write_markdown,
    ".txt": write_text,
    ".docx": write_docx,
    ".xlsx": write_xlsx,
    ".pdf": write_pdf,
}


def _is_filesystem_root(path: Path) -> bool:
    resolved = path.resolve()
    return resolved == Path(resolved.anchor)


def generate_demo_files(output_root: Path) -> tuple[Path, Path]:
    """在指定根目录下重建两棵演示目录，且不会删除其他路径。"""
    output_root = output_root.resolve()
    if _is_filesystem_root(output_root):
        raise ValueError("output_root 不能是文件系统根目录。")
    knowledge_root = output_root / KNOWLEDGE_DIRECTORY
    questions_root = output_root / QUESTIONS_DIRECTORY
    for directory in (knowledge_root, questions_root):
        if directory.exists():
            shutil.rmtree(directory)
    for spec in KNOWLEDGE_DOCUMENTS:
        target = knowledge_root / spec.folder / spec.filename
        target.parent.mkdir(parents=True, exist_ok=True)
        WRITERS[target.suffix](target, spec)
    for folder, questions in QUESTION_SETS.items():
        target = questions_root / folder / QUESTION_FILENAME
        target.parent.mkdir(parents=True, exist_ok=True)
        write_question_set(target, folder, questions)
    return knowledge_root, questions_root


def main() -> None:
    parser = argparse.ArgumentParser(description="生成演示知识库资料和测试问题。")
    parser.add_argument("--output-root", type=Path, required=True, help="项目根目录或其他安全输出目录")
    args = parser.parse_args()
    knowledge_root, questions_root = generate_demo_files(args.output_root)
    print(f"已生成知识资料：{knowledge_root}")
    print(f"已生成测试问题：{questions_root}")


if __name__ == "__main__":
    main()
