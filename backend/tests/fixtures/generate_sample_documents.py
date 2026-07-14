"""生成可上传到本项目的本地测试文档。"""

from pathlib import Path

import fitz
from docx import Document
from openpyxl import Workbook

DOCUMENTS_DIRECTORY = Path(__file__).parent / "documents"
MAX_UPLOAD_BYTES = 20 * 1024 * 1024


def write_text_documents() -> None:
    (DOCUMENTS_DIRECTORY / "01-年假制度.txt").write_text(
        "人事制度：员工入职满一年后，享有 5 天带薪年假。\n"
        "年假应在当年使用，特殊情况可由直属主管审批后顺延。\n",
        encoding="utf-8",
    )
    (DOCUMENTS_DIRECTORY / "02-信息安全规范.md").write_text(
        "# 信息安全规范\n\n"
        "## 密码要求\n\n"
        "密码长度不少于 12 位，并且每 90 天更新一次。\n\n"
        "## 设备管理\n\n"
        "离开工位超过 5 分钟时，必须锁定电脑屏幕。\n",
        encoding="utf-8",
    )
    duplicate_content = "这是一份用于验证重复文件检测的测试文本。\n"
    (DOCUMENTS_DIRECTORY / "06-重复内容-A.txt").write_text(duplicate_content, encoding="utf-8")
    (DOCUMENTS_DIRECTORY / "07-重复内容-B.txt").write_text(duplicate_content, encoding="utf-8")
    (DOCUMENTS_DIRECTORY / "08-空白内容.txt").write_text("\u00a0\n", encoding="utf-8")
    (DOCUMENTS_DIRECTORY / "09-不支持格式.csv").write_text(
        "姓名,部门\n张三,研发部\n", encoding="utf-8"
    )
    (DOCUMENTS_DIRECTORY / "10-损坏的PDF.pdf").write_bytes(b"not a valid pdf")
    (DOCUMENTS_DIRECTORY / "11-超过20MB限制.txt").write_bytes(b"A" * (MAX_UPLOAD_BYTES + 1))


def write_word_document() -> None:
    document = Document()
    document.add_heading("员工手册（测试版）", level=0)
    document.add_heading("考勤规则", level=1)
    document.add_paragraph("员工应在工作日 9:30 前完成打卡；迟到超过 30 分钟需提交补卡说明。")
    document.add_heading("请假规则", level=1)
    document.add_paragraph("病假申请应在当天 10:00 前通知直属主管，并在返岗后补交证明材料。")
    document.save(DOCUMENTS_DIRECTORY / "03-员工手册.docx")


def write_excel_document() -> None:
    workbook = Workbook()
    training = workbook.active
    training.title = "培训计划"
    training.append(["课程", "对象", "计划日期", "时长（小时）"])
    training.append(["信息安全基础", "全体员工", "2026-08-01", 2])
    training.append(["Python AI 入门", "研发部", "2026-08-15", 4])
    contacts = workbook.create_sheet("联系人")
    contacts.append(["部门", "联系人", "邮箱"])
    contacts.append(["人力资源部", "李老师", "hr@example.com"])
    workbook.save(DOCUMENTS_DIRECTORY / "04-培训计划.xlsx")


def write_pdf_document() -> None:
    pdf = fitz.open()
    font_name = "china-s"
    first_page = pdf.new_page()
    first_page.insert_text(
        (72, 72),
        "远程办公指南（测试版）\n\n员工每周最多可申请 2 天远程办公。",
        fontname=font_name,
    )
    second_page = pdf.new_page()
    second_page.insert_text(
        (72, 72),
        "远程办公期间\n\n员工必须通过公司 VPN 访问内部系统，并在工作时段保持在线。",
        fontname=font_name,
    )
    pdf.save(DOCUMENTS_DIRECTORY / "05-远程办公指南.pdf")
    pdf.close()


def main() -> None:
    DOCUMENTS_DIRECTORY.mkdir(parents=True, exist_ok=True)
    write_text_documents()
    write_word_document()
    write_excel_document()
    write_pdf_document()
    print(f"已生成测试文档：{DOCUMENTS_DIRECTORY}")


if __name__ == "__main__":
    main()
