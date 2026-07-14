from pathlib import Path

from docx import Document as WordDocument

from app.core.exceptions import AppError
from app.knowledge.schemas import ParsedSection


class WordParser:
    def parse(self, file_path: Path) -> list[ParsedSection]:
        document = WordDocument(file_path)
        sections = [
            ParsedSection(text=paragraph.text.strip())
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        for table in document.tables:
            for row in table.rows:
                text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if text:
                    sections.append(ParsedSection(text=text))
        if not sections:
            raise AppError(code="DOCUMENT_CONTENT_EMPTY", message="文档内容为空。", status_code=422)
        return sections
